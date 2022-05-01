from selenium import webdriver
from selenium.webdriver.common.by import By
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

file_name = input('请输入输出文件的名称:')
print('正在打开光明网时评网站...')

browser = webdriver.Chrome()
doc = Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.styles['Normal'].font.size = Pt(10.5)

articles = []

while True:
    browser.get('https://guancha.gmw.cn/node_26275.htm')
    result = []
    l = browser.find_element(By.XPATH, '/html/body/div[5]/div[1]/div[2]/ul').find_elements(By.TAG_NAME, 'li')

    i = 1
    for e in l:
        e2 = e.find_element(By.TAG_NAME, 'span').find_element(By.CSS_SELECTOR,'span.channel-newsTitle > a')
        print(f'{i}.' + e2.text)
        result.append(e2)
        i += 1

    #TODO 增加翻页功能
    page = 1
    n = input(f'以上是第 {page} 页内新闻。你想要哪篇文章（如已完成请输入end）:')
    if n == 'end':
        break
    browser.get(result[int(n)-1].get_attribute('href'))

    article = {}
    article['title'] = browser.find_element(By.CSS_SELECTOR,'body > div.g-main > div.m-title-box > h1').text
    temp = browser.find_element(By.XPATH, '//*[@id="article_inbox"]/div[6]').find_elements(By.TAG_NAME, 'p')
    text = []
    for p in temp:
        text.append(p.text)
    article['body'] = text[1:]
    articles.append(article)

for article in articles:
    doc.add_paragraph(article['title'])
    doc.paragraphs[-1].runs[0].font.size = Pt(15)
    doc.paragraphs[-1].runs[0].bold = True

    for p in article['body']:
        doc.add_paragraph('    ' + p)
        doc.paragraphs[-1].runs[0].size = Pt(10.5)

doc.save(file_name)
